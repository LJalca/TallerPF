from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Entregable_Taller_Programacion_Funcional_Java_Grupo_4.docx"
CAPTURAS = ROOT / "capturas"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "EAF3F8"
GRAY = "5B6573"
BLACK = "000000"


def set_font(run, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    code_style = doc.styles.add_style("Código", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(8.2)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = 1.0
    code_style.paragraph_format.left_indent = Inches(0.12)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Taller Funcional - Grupo 4 | Página ")
    set_font(r, size=8.5, color=GRAY)
    add_page_number(p)


def para(doc, text="", style=None, align=None, before=None, after=None, size=None, bold=None, color=None, italic=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=10.7)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=10.7)
    return p


def add_code(doc, code):
    p = doc.add_paragraph(style="Código")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(code)
    set_font(r, name="Consolas", size=8.2, color="1E1E1E")
    return p


def caption(doc, text):
    p = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=9)
    r = p.add_run(text)
    set_font(r, size=9, italic=True, color=GRAY)


def add_image(doc, filename, label, width=5.95):
    path = CAPTURAS / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(path), width=Inches(width))
        caption(doc, label)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_fill(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=110, start=150, bottom=110, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_font(r2, size=10.3)
    para(doc, "", after=2)


def cover(doc):
    para(doc, "", before=54, after=0)
    para(doc, "TRABAJO EN CLASE", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True, color=BLUE, after=12)
    para(doc, "Taller Funcional", align=WD_ALIGN_PARAGRAPH.CENTER, size=28, bold=True, color=DARK_BLUE, after=6)
    para(doc, "Programación Funcional en Java", align=WD_ALIGN_PARAGRAPH.CENTER, size=17, color=GRAY, after=44)
    para(doc, "Grupo: 4", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True, after=38)
    para(doc, "INTEGRANTES", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, bold=True, color=BLUE, after=8)
    for name in ["Jalca Saltos Luiggi", "Naranjo Christian", "Narváez Lopez Jilson"]:
        para(doc, name, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=4)
    para(doc, "", after=30)
    para(doc, "Materia", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, bold=True, color=BLUE, after=2)
    para(doc, "Desarrollo de aplicaciones empresariales", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=16)
    para(doc, "Docente", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, bold=True, color=BLUE, after=2)
    para(doc, "Ing. Christian Merchán Millán", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=30)
    para(doc, "Agosto de 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=GRAY)
    doc.add_page_break()


def section_overview(doc):
    para(doc, "1. Presentación", style="Heading 1")
    para(doc, "Este documento presenta el primer entregable del Taller de Programación Funcional en Java. El trabajo desarrolla una aplicación de consola que procesa una lista inmutable de empleados mediante expresiones lambda, interfaces funcionales, Streams, referencias a métodos y Optional.")
    add_callout(doc, "Resultado del entregable", "La clase TallerFuncional.java compila y ejecuta correctamente con JDK 17+. La solución cumple los siete pasos solicitados y añade el reto opcional de agrupación por rango salarial.")
    para(doc, "2. Objetivos del taller", style="Heading 1")
    add_bullet(doc, "Aplicar los principios de programación funcional en Java a partir de una lista de empleados.")
    add_bullet(doc, "Filtrar, transformar, ordenar y resumir datos utilizando la API Stream.")
    add_bullet(doc, "Usar Predicate, Function, referencias a métodos y Optional de forma segura.")
    add_bullet(doc, "Comprobar que la salida producida coincide con los resultados esperados del taller.")
    para(doc, "3. Entorno y ejecución", style="Heading 1")
    para(doc, "El proyecto se ejecutó con Java 17 o superior. Desde la carpeta del proyecto se emplean los siguientes comandos:")
    add_code(doc, "cd taller-funcional\njavac TallerFuncional.java\njava TallerFuncional")
    add_image(doc, "00-java-version.png", "Figura 1. Verificación de la versión de Java instalada.", 5.7)
    para(doc, "La estructura principal contiene el archivo fuente TallerFuncional.java y una carpeta de capturas que documenta la ejecución de cada paso.")
    doc.add_page_break()


def section_development_a(doc):
    para(doc, "4. Desarrollo del taller", style="Heading 1")
    para(doc, "4.1 Paso 1: clase base y comprobación del entorno", style="Heading 2")
    para(doc, "Se creó la clase pública TallerFuncional con el método main. Al iniciar, el programa confirma que el entorno está listo mediante un mensaje de consola.")
    add_code(doc, 'public static void main(String[] args) {\n    System.out.println("Taller de Programación Funcional iniciado");\n}')
    add_image(doc, "01-paso1.png", "Figura 2. Creación de la clase base y ejecución inicial.", 5.65)
    para(doc, "4.2 Paso 2: modelo Empleado y datos de prueba", style="Heading 2")
    para(doc, "Se definió un record inmutable para representar a cada empleado y se cargó una lista inmutable con cinco registros de prueba. Este enfoque evita modificaciones accidentales y se alinea con el paradigma funcional.")
    add_code(doc, 'record Empleado(String nombre, double salario) {}\n\nList<Empleado> empleados = List.of(\n    new Empleado("Carlos", 2500.0), new Empleado("Maria", 1800.0),\n    new Empleado("Pedro", 1200.0), new Empleado("Ana", 900.0),\n    new Empleado("Luis", 750.0)\n);')
    add_image(doc, "02-paso2.png", "Figura 3. Definición del modelo Empleado y lista de datos de prueba.", 5.65)
    doc.add_page_break()


def section_development_b(doc):
    para(doc, "4.3 Paso 3: filtrado con Predicate y Stream.filter", style="Heading 2")
    para(doc, "Se declaró un Predicate que selecciona empleados cuyo salario supera 1000. La operación filter conserva a Carlos, Maria y Pedro, por lo que el resultado contiene tres empleados.")
    add_code(doc, "Predicate<Empleado> salarioAlto = e -> e.salario() > 1000;\nList<Empleado> filtrados = empleados.stream()\n    .filter(salarioAlto)\n    .collect(Collectors.toList());")
    add_image(doc, "03-paso3.png", "Figura 4. Filtrado de empleados con salario mayor a 1000.", 5.85)
    para(doc, "4.4 Paso 4: Function, referencias a método y ordenamiento", style="Heading 2")
    para(doc, "La Function convierte los nombres a mayúsculas; la referencia Empleado::salario se utiliza al comparar y ordenar de forma descendente. Finalmente, forEach imprime el reporte de los empleados filtrados.")
    add_code(doc, "Function<Empleado, String> aMayusculas = e -> e.nombre().toUpperCase();\nList<Empleado> ordenados = filtrados.stream()\n    .sorted(Comparator.comparingDouble(Empleado::salario).reversed())\n    .collect(Collectors.toList());")
    add_image(doc, "04-paso4.png", "Figura 5. Transformación, referencia a método y ordenamiento por salario.", 5.85)
    doc.add_page_break()


def section_development_c(doc):
    para(doc, "4.5 Paso 5: totales con reduce y Collectors", style="Heading 2")
    para(doc, "Se calculó el salario total mediante map y reduce con Double::sum, y el promedio con Collectors.averagingDouble. Para los empleados filtrados, el total es 5500.0 y el promedio es 1833.33.")
    add_code(doc, "double total = ordenados.stream().map(Empleado::salario)\n    .reduce(0.0, Double::sum);\ndouble promedio = ordenados.stream()\n    .collect(Collectors.averagingDouble(Empleado::salario));")
    add_image(doc, "05-paso5.png", "Figura 6. Cálculo del total y promedio salarial.", 5.85)
    para(doc, "4.6 Paso 6: máximo con Optional", style="Heading 2")
    para(doc, "El empleado mejor pagado se obtiene mediante Stream.max, cuyo resultado es Optional<Empleado>. La resolución con ifPresentOrElse evita usar get() sin verificar la existencia de un valor.")
    add_code(doc, "Optional<Empleado> mejorPagado = ordenados.stream()\n    .max(Comparator.comparingDouble(Empleado::salario));\nmejorPagado.ifPresentOrElse(\n    e -> System.out.println(e.nombre().toUpperCase()),\n    () -> System.out.println(\"No hay empleados que cumplan el criterio\")\n);")
    add_image(doc, "06-paso6.png", "Figura 7. Obtención segura del empleado mejor pagado con Optional.", 5.85)
    doc.add_page_break()


def section_results(doc):
    para(doc, "5. Resultados y verificación final", style="Heading 1")
    para(doc, "La ejecución final confirma el funcionamiento de todos los componentes requeridos. El reporte muestra los empleados ordenados, el total, el promedio y el empleado mejor pagado.")
    add_image(doc, "07-entrega-final.png", "Figura 8. Ejecución final del programa y reto opcional.", 5.8)
    para(doc, "Resultados observados", style="Heading 2")
    add_bullet(doc, "Empleados con salario superior a 1000: Carlos, Maria y Pedro.")
    add_bullet(doc, "Orden descendente por salario: CARLOS (2500.0), MARIA (1800.0) y PEDRO (1200.0).")
    add_bullet(doc, "Salario total del equipo filtrado: 5500.0.")
    add_bullet(doc, "Salario promedio: 1833.33.")
    add_bullet(doc, "Empleado mejor pagado: CARLOS.")
    para(doc, "Reto opcional implementado", style="Heading 2")
    para(doc, "Además de lo solicitado, se implementó una agrupación de todos los empleados por rango salarial usando Collectors.groupingBy y se generó la lista de nombres de cada rango con Collectors.joining. Los rangos mostrados son ALTO, MEDIO y BAJO.")
    add_callout(doc, "Criterio funcional evidenciado", "La solución usa lambdas, referencias a método, Streams y Optional; además, conserva la lista original inmutable y genera nuevos resultados para cada operación.")
    doc.add_page_break()


def section_evaluation(doc):
    para(doc, "6. Autoevaluación", style="Heading 1")
    para(doc, "¿Qué ventaja ofrece usar Stream frente a un bucle for tradicional en este ejercicio?", style="Heading 2")
    para(doc, "Stream permite expresar qué se desea hacer -filtrar, ordenar o sumar- mediante una cadena declarativa. Así se evita mezclar la lógica de acumulación, comparación e impresión en un mismo bucle; el código es más breve, legible y se presta a paralelización si fuese necesaria.")
    para(doc, "¿Por qué Optional evita un error en tiempo de ejecución al buscar el empleado mejor pagado?", style="Heading 2")
    para(doc, "El método max() devuelve un Optional que puede estar vacío cuando no existen elementos. Con ifPresentOrElse se utiliza el empleado solamente cuando está presente y, de no haber coincidencias, se muestra un mensaje alternativo. Esto previene NullPointerException y evita llamar a get() sin comprobación previa.")
    para(doc, "¿Qué diferencia existe entre una expresión lambda y una referencia a método?", style="Heading 2")
    para(doc, "Una lambda define una función anónima, por ejemplo e -> e.salario() > 1000. Una referencia a método, como Empleado::salario o Double::sum, apunta a un método existente cuya firma coincide con la interfaz funcional requerida; por ello es una forma más concisa de expresar una lambda equivalente.")
    para(doc, "7. Conclusiones", style="Heading 1")
    para(doc, "El taller permitió aplicar de forma práctica las bases de la programación funcional en Java. La solución satisface los resultados esperados al procesar datos sin mutar la colección original y emplear operaciones funcionales para filtrar, transformar, ordenar y resumir la información. El uso de Optional fortalece el manejo seguro de resultados potencialmente ausentes.")
    doc.add_page_break()


def section_appendix(doc):
    para(doc, "Anexo A. Código fuente entregado", style="Heading 1")
    para(doc, "Archivo: taller-funcional/TallerFuncional.java", italic=True, color=GRAY, after=8)
    lines = (ROOT / "taller-funcional" / "TallerFuncional.java").read_text(encoding="utf-8").splitlines()
    add_code_block(doc, lines[:35])
    doc.add_page_break()
    para(doc, "Anexo A. Código fuente (continuación)", style="Heading 1")
    para(doc, "Archivo: taller-funcional/TallerFuncional.java", italic=True, color=GRAY, after=8)
    add_code_block(doc, lines[35:])


def add_code_block(doc, lines):
    p = doc.add_paragraph(style="Código")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    set_font(run, name="Consolas", size=8.2, color="1E1E1E")
    for index, line in enumerate(lines):
        run.add_text(line if line else " ")
        if index < len(lines) - 1:
            run.add_break()


def main():
    doc = Document()
    configure(doc)
    cover(doc)
    section_overview(doc)
    section_development_a(doc)
    section_development_b(doc)
    section_development_c(doc)
    section_results(doc)
    section_evaluation(doc)
    section_appendix(doc)
    doc.core_properties.title = "Entregable - Taller de Programación Funcional en Java"
    doc.core_properties.author = "Grupo 4"
    doc.core_properties.subject = "Desarrollo de aplicaciones empresariales"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
